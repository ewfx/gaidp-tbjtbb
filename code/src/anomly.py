import pandas as pd
import numpy as np
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from sklearn.ensemble import IsolationForest
from docx import Document

# Load a small open-source LLM
MODEL_NAME = "google/flan-t5-small"

# Load tokenizer and model
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME, torch_dtype=torch.float32)

def load_data(file_path):
    """Load dataset from CSV or Excel file."""
    try:
        if file_path.endswith(".csv"):
            df = pd.read_csv(file_path)
        elif file_path.endswith(".xlsx"):
            df = pd.read_excel(file_path)
        else:
            raise ValueError("Unsupported file format. Use CSV or Excel.")
        
        print(f"✅ Data loaded successfully. {df.shape[0]} rows, {df.shape[1]} columns.")
        return df
    except Exception as e:
        print(f"❌ Error loading file: {e}")
        return None

def generate_rules():
    """Use LLM to generate and explain structured validation rules."""
    prompt = """Generate financial validation rules in JSON format with descriptions:
    - Total_Assets and Liabilities must be positive numbers.
    - Net_Profit can be negative but should not exceed -1M.
    - Regulatory_Compliance must be 'Compliant' or 'Non-Compliant'.
    Also, provide a brief explanation of why these rules exist.
    Return structured rules only."""

    # Tokenize input
    inputs = tokenizer(prompt, return_tensors="pt").input_ids.to("cpu")

    # Generate text
    with torch.no_grad():
        output = model.generate(inputs, max_length=300)

    # Decode response
    rules = tokenizer.decode(output[0], skip_special_tokens=True)
    print("📜 Generated Rules:\n", rules)
    return rules

def validate_data(df):
    """Perform rule-based validation and detect anomalies."""
    errors = []

    # Check for missing values
    for column in df.columns:
        missing_count = df[column].isnull().sum()
        if missing_count > 0:
            errors.append(f"⚠️ Column '{column}' has {missing_count} missing values.")

    # Check for negative financial values (except Net_Profit)
    for column in ["Total_Assets", "Liabilities"]:
        if column in df.columns and (df[column] < 0).any():
            errors.append(f"❌ Column '{column}' contains negative values, which is invalid.")

    # Check for extreme Net_Profit loss
    if "Net_Profit" in df.columns and (df["Net_Profit"] < -1_000_000).any():
        errors.append(f"⚠️ Some 'Net_Profit' values exceed the allowed loss of -1M.")

    # Validate compliance status
    if "Regulatory_Compliance" in df.columns:
        valid_status = ["Compliant", "Non-Compliant"]
        invalid_entries = df[~df["Regulatory_Compliance"].isin(valid_status)]
        if not invalid_entries.empty:
            errors.append(f"❌ Invalid values detected in 'Regulatory_Compliance' column: {invalid_entries['Regulatory_Compliance'].unique()}")

    return errors

def detect_anomalies(df):
    """Use Isolation Forest to detect anomalies and describe them."""
    anomalies = []

    numeric_cols = df.select_dtypes(include=[np.number]).columns
    if len(numeric_cols) == 0:
        return ["⚠️ No numeric columns found for anomaly detection."]

    contamination_rate = min(0.1, 5 / len(df))  # Adjust dynamically based on dataset size
    model = IsolationForest(contamination=contamination_rate, random_state=42)
    df["Anomaly"] = model.fit_predict(df[numeric_cols])

    anomaly_rows = df[df["Anomaly"] == -1]
    if not anomaly_rows.empty:
        for idx, row in anomaly_rows.iterrows():
            anomaly_details = ", ".join(
                [f"{col}: {row[col]}" for col in numeric_cols if row[col] > df[col].mean() + 2 * df[col].std() or row[col] < df[col].mean() - 2 * df[col].std()]
            )
            anomalies.append(f"🚨 Row {idx} is an anomaly. Suspicious values: {anomaly_details}")

    return anomalies

def generate_report(errors, anomalies, rules, report_path="validation_report.docx"):
    """Generate and save a detailed validation report in Word format."""
    doc = Document()
    doc.add_heading("📊 Financial Data Validation Report", level=1)

    doc.add_heading("📝 Generated Validation Rules", level=2)
    doc.add_paragraph(rules)

    doc.add_heading("❌ Basic Validation Errors", level=2)
    if errors:
        for error in errors:
            doc.add_paragraph(f"- {error}")
    else:
        doc.add_paragraph("✅ No rule-based validation errors detected.")

    doc.add_heading("🔍 AI-Driven Anomaly Detection", level=2)
    if anomalies:
        for anomaly in anomalies:
            doc.add_paragraph(f"- {anomaly}")
    else:
        doc.add_paragraph("✅ No anomalies detected.")

    doc.save(report_path)
    print(f"📄 Validation report saved to {report_path}")

if __name__ == "__main__":
    file_path = input("📂 Enter the path to the CSV/Excel file: ")
    df = load_data(file_path)

    if df is not None:
        rules = generate_rules()
        validation_errors = validate_data(df)
        anomaly_results = detect_anomalies(df)
        generate_report(validation_errors, anomaly_results, rules)
